from docx import Document
import re
from pathlib import Path

def extract_text_and_hyperlinks(docx_path):
    document = Document(docx_path)
    results = []

    for para in document.paragraphs:
        full_text = ""
        links = []

        for run in para.runs:
            full_text += run.text

        # Identify hyperlinks
        hyperlink_map = {}
        for rel in document.part.rels.values():
            if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink":
                hyperlink_map[rel.target_ref] = rel.target_ref

        # Extract links from XML if they exist in paragraph
        for hyperlink in para._element.xpath('.//w:hyperlink'):
            r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if r_id:
                link = document.part.rels[r_id].target_ref
                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                links.append((text, link))

        # Split full_text into sentences
        sentences = re.split(r'(?<=[.!?]) +', full_text)
        for sentence in sentences:
            sentence_links = [(text, link) for text, link in links if text in sentence]
            results.append((sentence, sentence_links))

    return results

def generate_html(sentences_with_links, output_file):
    html = ['<html><body>']
    for sentence, links in sentences_with_links:
        html.append('<p>')
        html.append(sentence)
        if links:
            html.append('<ul>')
            for text, url in links:
                html.append(f'<li><a href="{url}" target="_blank">{text}</a></li>')
            html.append('</ul>')
        html.append('</p>')
    html.append('</body></html>')

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html))

if __name__ == "__main__":
    docx_input = "input.docx"
    html_output = "output.html"
    
    extracted = extract_text_and_hyperlinks(docx_input)
    generate_html(extracted, html_output)
    print(f"[+] Output saved to {html_output}")








from docx import Document
import re

def extract_sentences_with_links(docx_path, output_path="output.txt"):
    doc = Document(docx_path)
    results = []

    for para in doc.paragraphs:
        full_text = ''
        links = []

        for run in para.runs:
            full_text += run.text

        # Extract hyperlinks from paragraph XML
        para_xml = para._element.xml
        hyperlink_matches = re.findall(r'w:hyperlink[^>]*r:id="([^"]+)"', para_xml)
        hyperlink_targets = []

        for rId in hyperlink_matches:
            try:
                hyperlink = doc.part.related_parts[doc.part.rels[rId].target_ref]
                hyperlink_targets.append(hyperlink)
            except:
                hyperlink_targets.append("N/A")

        # Sentence split
        sentences = re.split(r'(?<=[.!?])\s+', full_text)
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            found_link = None
            for target in hyperlink_targets:
                if isinstance(target, str) and target in sentence:
                    found_link = target
                    break
                elif hasattr(target, 'target_ref') and target.target_ref in sentence:
                    found_link = target.target_ref
                    break

            results.append((sentence, found_link if found_link else "N/A"))

    # Save to file
    with open(output_path, "w", encoding="utf-8") as f:
        for sentence, link in results:
            f.write(f"Sentence: {sentence}\nLink: {link}\n\n")

# Example usage
extract_sentences_with_links("input.docx")
